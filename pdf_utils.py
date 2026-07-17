"""
===========================================================
Universal PDF Merger & Viewer
Utility Functions
Author : PragyanAI
===========================================================
"""

import os
import io
import uuid
import shutil
import tempfile
from pathlib import Path
from typing import List, Tuple

from PIL import Image

import fitz  # PyMuPDF
from pypdf import PdfReader, PdfWriter, PdfMerger

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader

from docx import Document

###########################################################
# CONFIGURATION
###########################################################

TEMP_FOLDER = "temp"
OUTPUT_FOLDER = "output"

os.makedirs(TEMP_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

SUPPORTED_IMAGE_TYPES = {
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tiff",
    ".tif",
    ".gif",
}

SUPPORTED_DOC_TYPES = {
    ".docx",
}

SUPPORTED_PDF_TYPES = {
    ".pdf",
}

SUPPORTED_TYPES = (
    SUPPORTED_IMAGE_TYPES
    | SUPPORTED_DOC_TYPES
    | SUPPORTED_PDF_TYPES
)

###########################################################
# CREATE TEMP DIRECTORY
###########################################################

def create_temp_directory():
    """
    Creates a unique temporary directory.

    Returns
    -------
    str
        Temporary directory path
    """

    folder = tempfile.mkdtemp(prefix="pdfmerge_")
    return folder


###########################################################
# CLEAN TEMP DIRECTORY
###########################################################

def delete_directory(folder):

    if os.path.exists(folder):
        shutil.rmtree(folder, ignore_errors=True)


###########################################################
# RANDOM FILE NAME
###########################################################

def random_filename(extension=".pdf"):
    """
    Returns

    87ab8898a99b.pdf
    """

    return f"{uuid.uuid4().hex}{extension}"


###########################################################
# SAVE STREAMLIT UPLOADED FILE
###########################################################

def save_uploaded_file(uploaded_file, folder=TEMP_FOLDER):
    """
    Saves UploadedFile object to disk.

    Parameters
    ----------
    uploaded_file

    Returns
    -------
    path
    """

    extension = Path(uploaded_file.name).suffix.lower()

    filename = random_filename(extension)

    save_path = os.path.join(folder, filename)

    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    return save_path


###########################################################
# FILE EXTENSION
###########################################################

def get_extension(file_path):

    return Path(file_path).suffix.lower()


###########################################################
# FILE TYPE
###########################################################

def get_file_type(file_path):

    ext = get_extension(file_path)

    if ext in SUPPORTED_PDF_TYPES:
        return "pdf"

    if ext in SUPPORTED_DOC_TYPES:
        return "doc"

    if ext in SUPPORTED_IMAGE_TYPES:
        return "image"

    return "unknown"


###########################################################
# VALIDATE FILE
###########################################################

def is_supported(file_path):

    ext = get_extension(file_path)

    return ext in SUPPORTED_TYPES


###########################################################
# FILE SIZE
###########################################################

def readable_size(file_path):

    size = os.path.getsize(file_path)

    units = ["B", "KB", "MB", "GB"]

    i = 0

    while size > 1024 and i < len(units)-1:
        size /= 1024
        i += 1

    return f"{size:.2f} {units[i]}"


###########################################################
# FILE EXISTS
###########################################################

def file_exists(path):

    return os.path.exists(path)


###########################################################
# PDF VALIDATION
###########################################################

def validate_pdf(pdf_path):
    """
    Checks whether PDF can be opened.

    Returns
    -------
    bool
    """

    try:

        PdfReader(pdf_path)

        return True

    except Exception:

        return False


###########################################################
# IMAGE VALIDATION
###########################################################

def validate_image(image_path):

    try:

        img = Image.open(image_path)

        img.verify()

        return True

    except Exception:

        return False


###########################################################
# DOCX VALIDATION
###########################################################

def validate_docx(doc_path):

    try:

        Document(doc_path)

        return True

    except Exception:

        return False


###########################################################
# LOGGING
###########################################################

def log(message):

    print(f"[INFO] {message}")


def warning(message):

    print(f"[WARNING] {message}")


def error(message):

    print(f"[ERROR] {message}")


###########################################################
# FILE INFORMATION
###########################################################

def file_information(file_path):

    return {
        "name": os.path.basename(file_path),
        "path": file_path,
        "type": get_file_type(file_path),
        "size": readable_size(file_path),
        "exists": file_exists(file_path)
    }


###########################################################
# LIST OF FILE INFORMATION
###########################################################

def list_files_information(files):

    return [file_information(f) for f in files]
###########################################################
# IMAGE TO PDF
###########################################################

from PIL import ImageOps

A4_WIDTH, A4_HEIGHT = A4


def fix_image_rotation(image):
    """
    Automatically applies EXIF orientation.
    """

    try:
        image = ImageOps.exif_transpose(image)
    except Exception:
        pass

    return image


###########################################################
# CONVERT IMAGE TO RGB
###########################################################

def image_to_rgb(image):
    """
    Converts image into RGB.

    Handles transparency.
    """

    if image.mode == "RGBA":

        background = Image.new("RGB", image.size, (255, 255, 255))
        background.paste(image, mask=image.split()[3])

        return background

    elif image.mode == "LA":

        background = Image.new("RGB", image.size, (255, 255, 255))
        background.paste(image)

        return background

    elif image.mode != "RGB":

        return image.convert("RGB")

    return image


###########################################################
# RESIZE IMAGE FOR A4
###########################################################

def resize_for_a4(image):
    """
    Returns resized dimensions that fit within
    an A4 page while maintaining aspect ratio.
    """

    width, height = image.size

    max_width = A4_WIDTH - 40
    max_height = A4_HEIGHT - 40

    ratio = min(max_width / width, max_height / height)

    new_width = width * ratio
    new_height = height * ratio

    return new_width, new_height


###########################################################
# DRAW IMAGE ON PDF PAGE
###########################################################

def draw_image(canvas_obj, image):
    """
    Draw image centered on A4 page.
    """

    image = fix_image_rotation(image)
    image = image_to_rgb(image)

    width, height = resize_for_a4(image)

    x = (A4_WIDTH - width) / 2
    y = (A4_HEIGHT - height) / 2

    image_reader = ImageReader(image)

    canvas_obj.drawImage(
        image_reader,
        x,
        y,
        width=width,
        height=height,
        preserveAspectRatio=True,
    )


###########################################################
# SINGLE IMAGE -> PDF
###########################################################

def image_to_pdf(image_path, output_pdf=None):
    """
    Converts one image into one PDF.
    """

    if output_pdf is None:
        output_pdf = os.path.join(
            TEMP_FOLDER,
            random_filename(".pdf")
        )

    image = Image.open(image_path)

    c = canvas.Canvas(output_pdf, pagesize=A4)

    draw_image(c, image)

    c.showPage()
    c.save()

    return output_pdf


###########################################################
# MULTIPLE IMAGES -> PDF
###########################################################

def images_to_pdf(image_paths, output_pdf=None):
    """
    Merge multiple images into one PDF.
    """

    if output_pdf is None:
        output_pdf = os.path.join(
            TEMP_FOLDER,
            random_filename(".pdf")
        )

    c = canvas.Canvas(output_pdf, pagesize=A4)

    for image_path in image_paths:

        try:

            image = Image.open(image_path)

            draw_image(c, image)

            c.showPage()

        except Exception as e:

            warning(f"Cannot process image: {image_path}")
            warning(str(e))

    c.save()

    return output_pdf


###########################################################
# IMAGE DIMENSIONS
###########################################################

def get_image_size(image_path):

    image = Image.open(image_path)

    return image.size


###########################################################
# IMAGE INFORMATION
###########################################################

def image_information(image_path):

    image = Image.open(image_path)

    return {

        "width": image.width,

        "height": image.height,

        "mode": image.mode,

        "format": image.format,

        "size": readable_size(image_path)

    }


###########################################################
# VERIFY IMAGE LIST
###########################################################

def verify_images(image_paths):
    """
    Returns only valid images.
    """

    valid = []

    for image in image_paths:

        if validate_image(image):

            valid.append(image)

        else:

            warning(f"Invalid image skipped: {image}")

    return valid


###########################################################
# CONVERT UPLOADED IMAGE FILES
###########################################################

def convert_uploaded_images(uploaded_paths):
    """
    Converts each uploaded image into a PDF.

    Returns
    -------
    List[str]
        List of generated PDF paths.
    """

    pdf_files = []

    for image_path in uploaded_paths:

        pdf = image_to_pdf(image_path)

        pdf_files.append(pdf)

    return pdf_files
###########################################################
# DOCX -> PDF CONVERTER
###########################################################

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle
)

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT

styles = getSampleStyleSheet()

NORMAL_STYLE = styles["BodyText"]
NORMAL_STYLE.alignment = TA_LEFT
NORMAL_STYLE.spaceAfter = 8

HEADING1_STYLE = styles["Heading1"]
HEADING2_STYLE = styles["Heading2"]
HEADING3_STYLE = styles["Heading3"]


###########################################################
# PARAGRAPH STYLE
###########################################################

def get_paragraph_style(style_name):

    if not style_name:
        return NORMAL_STYLE

    style_name = style_name.lower()

    if "heading 1" in style_name:
        return HEADING1_STYLE

    if "heading 2" in style_name:
        return HEADING2_STYLE

    if "heading 3" in style_name:
        return HEADING3_STYLE

    return NORMAL_STYLE


###########################################################
# PARAGRAPH -> REPORTLAB
###########################################################

def paragraph_to_story(paragraph):

    story = []

    text = paragraph.text.strip()

    if not text:

        story.append(Spacer(1, 8))
        return story

    style = get_paragraph_style(paragraph.style.name)

    story.append(
        Paragraph(text.replace("\n", "<br/>"), style)
    )

    return story


###########################################################
# TABLE -> REPORTLAB
###########################################################

def table_to_story(table):

    data = []

    for row in table.rows:

        values = []

        for cell in row.cells:

            values.append(cell.text)

        data.append(values)

    tbl = Table(data)

    tbl.setStyle(

        TableStyle([

            ("GRID", (0,0), (-1,-1), 0.5, colors.black),

            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),

            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),

            ("BOTTOMPADDING", (0,0), (-1,0), 8),

            ("TOPPADDING", (0,0), (-1,-1), 6),

            ("BOTTOMPADDING", (0,0), (-1,-1), 6),

            ("LEFTPADDING", (0,0), (-1,-1), 6),

            ("RIGHTPADDING", (0,0), (-1,-1), 6),

        ])

    )

    return tbl


###########################################################
# EXTRACT STORY FROM DOCX
###########################################################

def extract_doc_story(doc):

    story = []

    ###################################################
    # Paragraphs
    ###################################################

    for paragraph in doc.paragraphs:

        story.extend(paragraph_to_story(paragraph))

    ###################################################
    # Tables
    ###################################################

    if len(doc.tables):

        story.append(Spacer(1, 12))

    for tbl in doc.tables:

        story.append(table_to_story(tbl))

        story.append(Spacer(1, 12))

    return story


###########################################################
# DOCX -> PDF
###########################################################

def docx_to_pdf(docx_path, output_pdf=None):

    if output_pdf is None:

        output_pdf = os.path.join(
            TEMP_FOLDER,
            random_filename(".pdf")
        )

    document = Document(docx_path)

    story = extract_doc_story(document)

    pdf = SimpleDocTemplate(

        output_pdf,

        pagesize=A4,

        leftMargin=40,

        rightMargin=40,

        topMargin=40,

        bottomMargin=40

    )

    pdf.build(story)

    return output_pdf


###########################################################
# MULTIPLE DOCX -> PDF
###########################################################

def convert_uploaded_docx(docx_files):

    pdfs = []

    for doc in docx_files:

        try:

            pdf = docx_to_pdf(doc)

            pdfs.append(pdf)

        except Exception as e:

            warning(f"Cannot convert DOCX : {doc}")

            warning(str(e))

    return pdfs


###########################################################
# DOCX INFORMATION
###########################################################

def docx_information(docx_path):

    document = Document(docx_path)

    return {

        "paragraphs": len(document.paragraphs),

        "tables": len(document.tables),

        "sections": len(document.sections),

        "size": readable_size(docx_path)

    }


###########################################################
# VERIFY DOCX LIST
###########################################################

def verify_docx(doc_list):

    valid = []

    for doc in doc_list:

        if validate_docx(doc):

            valid.append(doc)

        else:

            warning(f"Invalid DOCX skipped : {doc}")

    return valid
  ###########################################################
# PDF MERGING
###########################################################

from pypdf import PdfMerger, PdfReader
import os


###########################################################
# PDF PAGE COUNT
###########################################################

def pdf_page_count(pdf_path):
    """
    Returns the number of pages in a PDF.
    """

    try:

        reader = PdfReader(pdf_path)

        return len(reader.pages)

    except Exception as e:

        warning(f"Cannot read page count: {pdf_path}")
        warning(str(e))

        return 0


###########################################################
# TOTAL PAGES
###########################################################

def total_pages(pdf_list):
    """
    Total pages across multiple PDFs.
    """

    pages = 0

    for pdf in pdf_list:

        pages += pdf_page_count(pdf)

    return pages


###########################################################
# VERIFY PDF LIST
###########################################################

def verify_pdfs(pdf_list):
    """
    Returns only valid PDFs.
    """

    valid = []

    for pdf in pdf_list:

        if validate_pdf(pdf):

            valid.append(pdf)

        else:

            warning(f"Invalid PDF skipped : {pdf}")

    return valid


###########################################################
# MERGE PDF LIST
###########################################################

def merge_pdfs(pdf_files, output_pdf=None):
    """
    Merge multiple PDFs into a single PDF.

    Parameters
    ----------
    pdf_files : list

    output_pdf : str

    Returns
    -------
    output pdf path
    """

    if output_pdf is None:

        output_pdf = os.path.join(
            OUTPUT_FOLDER,
            random_filename(".pdf")
        )

    pdf_files = verify_pdfs(pdf_files)

    if len(pdf_files) == 0:

        raise Exception("No valid PDFs to merge.")

    merger = PdfMerger()

    try:

        for pdf in pdf_files:

            log(f"Appending : {pdf}")

            merger.append(pdf)

        merger.write(output_pdf)

        merger.close()

    except Exception as e:

        merger.close()

        raise e

    return output_pdf


###########################################################
# MERGE EVERYTHING
###########################################################

def merge_all_files(file_paths, output_pdf=None):
    """
    Accepts any supported file types.

    PDF
    DOCX
    Images

    Converts everything to PDF and merges.
    """

    pdf_files = []

    image_files = []

    doc_files = []

    ##################################################
    # Separate file types
    ##################################################

    for file in file_paths:

        file_type = get_file_type(file)

        if file_type == "pdf":

            pdf_files.append(file)

        elif file_type == "image":

            image_files.append(file)

        elif file_type == "doc":

            doc_files.append(file)

    ##################################################
    # Convert Images
    ##################################################

    if len(image_files):

        log("Converting Images...")

        image_pdf = images_to_pdf(image_files)

        pdf_files.append(image_pdf)

    ##################################################
    # Convert DOCX
    ##################################################

    if len(doc_files):

        log("Converting DOCX...")

        converted = convert_uploaded_docx(doc_files)

        pdf_files.extend(converted)

    ##################################################
    # Merge
    ##################################################

    merged_pdf = merge_pdfs(

        pdf_files,

        output_pdf

    )

    return merged_pdf


###########################################################
# PDF INFORMATION
###########################################################

def pdf_information(pdf_path):
    """
    Returns PDF metadata.
    """

    reader = PdfReader(pdf_path)

    metadata = reader.metadata

    return {

        "pages": len(reader.pages),

        "title": metadata.title if metadata else "",

        "author": metadata.author if metadata else "",

        "producer": metadata.producer if metadata else "",

        "creator": metadata.creator if metadata else "",

        "encrypted": reader.is_encrypted,

        "size": readable_size(pdf_path)

    }


###########################################################
# MERGE SUMMARY
###########################################################

def merge_summary(pdf_files):

    summary = []

    total = 0

    for pdf in pdf_files:

        pages = pdf_page_count(pdf)

        total += pages

        summary.append({

            "name": os.path.basename(pdf),

            "pages": pages,

            "size": readable_size(pdf)

        })

    return {

        "files": summary,

        "total_files": len(pdf_files),

        "total_pages": total

    }


###########################################################
# CHECK IF PDF IS ENCRYPTED
###########################################################

def is_pdf_encrypted(pdf_path):

    try:

        reader = PdfReader(pdf_path)

        return reader.is_encrypted

    except Exception:

        return False


###########################################################
# REMOVE BLANK PDF FILES
###########################################################

def remove_empty_pdfs(pdf_files):

    valid = []

    for pdf in pdf_files:

        if pdf_page_count(pdf) > 0:

            valid.append(pdf)

    return valid
###########################################################
# PDF VIEWER / PAGE RENDERER
###########################################################

import fitz
from PIL import Image
import io


###########################################################
# OPEN PDF
###########################################################

def open_pdf(pdf_path):
    """
    Opens a PDF using PyMuPDF.
    """
    try:
        return fitz.open(pdf_path)
    except Exception as e:
        warning(f"Cannot open PDF : {pdf_path}")
        warning(str(e))
        return None


###########################################################
# GET TOTAL PAGES
###########################################################

def get_pdf_pages(pdf_path):

    doc = open_pdf(pdf_path)

    if doc is None:
        return 0

    pages = doc.page_count

    doc.close()

    return pages


###########################################################
# GET PAGE OBJECT
###########################################################

def get_pdf_page(pdf_path, page_number):

    doc = open_pdf(pdf_path)

    if doc is None:
        return None

    if page_number < 0:
        doc.close()
        return None

    if page_number >= doc.page_count:
        doc.close()
        return None

    page = doc.load_page(page_number)

    return doc, page


###########################################################
# RENDER PAGE
###########################################################

def render_page(pdf_path, page_number=0, zoom=2.0):
    """
    Returns a PIL image.
    """

    result = get_pdf_page(pdf_path, page_number)

    if result is None:
        return None

    doc, page = result

    matrix = fitz.Matrix(zoom, zoom)

    pix = page.get_pixmap(matrix=matrix)

    image = Image.open(
        io.BytesIO(pix.tobytes("png"))
    )

    doc.close()

    return image


###########################################################
# RENDER PAGE AS BYTES
###########################################################

def render_page_bytes(pdf_path, page_number=0, zoom=2):

    img = render_page(pdf_path, page_number, zoom)

    if img is None:
        return None

    buffer = io.BytesIO()

    img.save(buffer, format="PNG")

    return buffer.getvalue()


###########################################################
# THUMBNAIL
###########################################################

def render_thumbnail(pdf_path, page_number=0):

    return render_page(

        pdf_path,

        page_number,

        zoom=0.5

    )


###########################################################
# ALL PAGE IMAGES
###########################################################

def render_all_pages(pdf_path, zoom=1.5):

    images = []

    total = get_pdf_pages(pdf_path)

    for page in range(total):

        img = render_page(

            pdf_path,

            page,

            zoom

        )

        images.append(img)

    return images


###########################################################
# PAGE SIZE
###########################################################

def get_page_size(pdf_path, page_number=0):

    result = get_pdf_page(

        pdf_path,

        page_number

    )

    if result is None:
        return None

    doc, page = result

    rect = page.rect

    width = rect.width

    height = rect.height

    doc.close()

    return width, height


###########################################################
# PAGE ROTATION
###########################################################

def get_page_rotation(pdf_path, page_number=0):

    result = get_pdf_page(

        pdf_path,

        page_number

    )

    if result is None:
        return 0

    doc, page = result

    rotation = page.rotation

    doc.close()

    return rotation


###########################################################
# PAGE TEXT
###########################################################

def extract_page_text(pdf_path, page_number=0):

    result = get_pdf_page(

        pdf_path,

        page_number

    )

    if result is None:
        return ""

    doc, page = result

    text = page.get_text()

    doc.close()

    return text


###########################################################
# PDF TEXT
###########################################################

def extract_pdf_text(pdf_path):

    text = ""

    doc = open_pdf(pdf_path)

    if doc is None:
        return text

    for page in doc:

        text += page.get_text()

        text += "\n"

    doc.close()

    return text


###########################################################
# PDF INFORMATION
###########################################################

def pdf_properties(pdf_path):

    doc = open_pdf(pdf_path)

    if doc is None:

        return {}

    metadata = doc.metadata

    info = {

        "pages": doc.page_count,

        "title": metadata.get("title", ""),

        "author": metadata.get("author", ""),

        "creator": metadata.get("creator", ""),

        "producer": metadata.get("producer", ""),

        "subject": metadata.get("subject", ""),

        "keywords": metadata.get("keywords", "")

    }

    doc.close()

    return info


###########################################################
# FIRST PAGE IMAGE
###########################################################

def first_page_image(pdf_path):

    return render_page(

        pdf_path,

        0,

        zoom=2

    )


###########################################################
# LAST PAGE IMAGE
###########################################################

def last_page_image(pdf_path):

    pages = get_pdf_pages(pdf_path)

    if pages == 0:
        return None

    return render_page(

        pdf_path,

        pages - 1,

        zoom=2

    )
  ###########################################################
# PDF NAVIGATION HELPERS
###########################################################

DEFAULT_ZOOM = 2.0
MIN_ZOOM = 0.5
MAX_ZOOM = 5.0
ZOOM_STEP = 0.25


###########################################################
# CREATE VIEWER STATE
###########################################################

def create_viewer_state(total_pages):
    """
    Returns initial viewer state.
    """

    return {

        "page": 0,

        "total_pages": total_pages,

        "zoom": DEFAULT_ZOOM,

        "rotation": 0

    }


###########################################################
# PAGE VALIDATION
###########################################################

def clamp_page(page, total_pages):

    if total_pages <= 0:
        return 0

    if page < 0:
        return 0

    if page >= total_pages:
        return total_pages - 1

    return page


###########################################################
# NEXT PAGE
###########################################################

def next_page(state):

    state["page"] = clamp_page(

        state["page"] + 1,

        state["total_pages"]

    )

    return state


###########################################################
# PREVIOUS PAGE
###########################################################

def previous_page(state):

    state["page"] = clamp_page(

        state["page"] - 1,

        state["total_pages"]

    )

    return state


###########################################################
# GO TO PAGE
###########################################################

def goto_page(state, page_number):

    state["page"] = clamp_page(

        page_number,

        state["total_pages"]

    )

    return state


###########################################################
# FIRST PAGE
###########################################################

def first_page(state):

    state["page"] = 0

    return state


###########################################################
# LAST PAGE
###########################################################

def last_page(state):

    if state["total_pages"] > 0:

        state["page"] = state["total_pages"] - 1

    return state


###########################################################
# ZOOM IN
###########################################################

def zoom_in(state):

    zoom = state["zoom"] + ZOOM_STEP

    if zoom > MAX_ZOOM:

        zoom = MAX_ZOOM

    state["zoom"] = zoom

    return state


###########################################################
# ZOOM OUT
###########################################################

def zoom_out(state):

    zoom = state["zoom"] - ZOOM_STEP

    if zoom < MIN_ZOOM:

        zoom = MIN_ZOOM

    state["zoom"] = zoom

    return state


###########################################################
# SET ZOOM
###########################################################

def set_zoom(state, zoom):

    zoom = max(MIN_ZOOM, min(MAX_ZOOM, zoom))

    state["zoom"] = zoom

    return state


###########################################################
# ROTATE PAGE
###########################################################

def rotate_left(state):

    state["rotation"] -= 90

    return state


def rotate_right(state):

    state["rotation"] += 90

    return state


###########################################################
# CURRENT PAGE IMAGE
###########################################################

def current_page_image(pdf_path, state):

    image = render_page(

        pdf_path,

        state["page"],

        state["zoom"]

    )

    if image is None:

        return None

    rotation = state.get("rotation", 0)

    if rotation != 0:

        image = image.rotate(

            rotation,

            expand=True

        )

    return image


###########################################################
# PAGE LABEL
###########################################################

def current_page_label(state):

    return f"Page {state['page'] + 1} / {state['total_pages']}"


###########################################################
# PAGE PERCENTAGE
###########################################################

def viewing_progress(state):

    if state["total_pages"] == 0:

        return 0

    return round(

        ((state["page"] + 1)

         / state["total_pages"]) * 100,

        1

    )


###########################################################
# THUMBNAILS
###########################################################

def generate_thumbnails(pdf_path):

    thumbs = []

    total = get_pdf_pages(pdf_path)

    for page in range(total):

        thumbs.append(

            render_thumbnail(

                pdf_path,

                page

            )

        )

    return thumbs


###########################################################
# CACHE PAGE IMAGES
###########################################################

def cache_page_images(pdf_path, zoom=1.5):

    cache = {}

    total = get_pdf_pages(pdf_path)

    for page in range(total):

        cache[page] = render_page(

            pdf_path,

            page,

            zoom

        )

    return cache


###########################################################
# PAGE SEARCH
###########################################################

def search_text(pdf_path, keyword):

    results = []

    keyword = keyword.lower()

    doc = open_pdf(pdf_path)

    if doc is None:

        return results

    for page_no, page in enumerate(doc):

        text = page.get_text().lower()

        if keyword in text:

            results.append(page_no)

    doc.close()

    return results


###########################################################
# PAGE EXISTS
###########################################################

def page_exists(pdf_path, page_number):

    total = get_pdf_pages(pdf_path)

    return 0 <= page_number < total


###########################################################
# DOCUMENT SUMMARY
###########################################################

def document_summary(pdf_path):

    return {

        "pages": get_pdf_pages(pdf_path),

        "properties": pdf_properties(pdf_path),

        "first_page": first_page_image(pdf_path),

        "last_page": last_page_image(pdf_path)

    }
  ###########################################################
# CLEANUP & ADVANCED UTILITIES
###########################################################

import hashlib
from datetime import datetime


###########################################################
# DELETE FILE
###########################################################

def delete_file(file_path):
    """
    Safely delete a file.
    """

    try:

        if os.path.exists(file_path):

            os.remove(file_path)

            return True

    except Exception as e:

        warning(str(e))

    return False


###########################################################
# DELETE FILE LIST
###########################################################

def delete_files(file_list):

    deleted = 0

    for file in file_list:

        if delete_file(file):

            deleted += 1

    return deleted


###########################################################
# CLEAN TEMP FOLDER
###########################################################

def clean_temp_folder():

    if not os.path.exists(TEMP_FOLDER):

        return

    for file in os.listdir(TEMP_FOLDER):

        path = os.path.join(

            TEMP_FOLDER,

            file

        )

        delete_file(path)


###########################################################
# CLEAN OUTPUT FOLDER
###########################################################

def clean_output_folder():

    if not os.path.exists(OUTPUT_FOLDER):

        return

    for file in os.listdir(OUTPUT_FOLDER):

        path = os.path.join(

            OUTPUT_FOLDER,

            file

        )

        delete_file(path)


###########################################################
# SHA256 HASH
###########################################################

def sha256(file_path):

    h = hashlib.sha256()

    with open(file_path, "rb") as f:

        while True:

            chunk = f.read(8192)

            if not chunk:

                break

            h.update(chunk)

    return h.hexdigest()


###########################################################
# REMOVE DUPLICATES
###########################################################

def remove_duplicate_files(files):

    unique = []

    hashes = set()

    for file in files:

        try:

            digest = sha256(file)

            if digest not in hashes:

                hashes.add(digest)

                unique.append(file)

        except Exception:

            unique.append(file)

    return unique


###########################################################
# SORT FILES
###########################################################

def sort_files(files):

    return sorted(

        files,

        key=lambda x: os.path.basename(x).lower()

    )


###########################################################
# SORT BY SIZE
###########################################################

def sort_files_by_size(files):

    return sorted(

        files,

        key=lambda x: os.path.getsize(x)

    )


###########################################################
# SORT BY DATE
###########################################################

def sort_files_by_date(files):

    return sorted(

        files,

        key=lambda x: os.path.getmtime(x)

    )


###########################################################
# TOTAL FILE SIZE
###########################################################

def total_file_size(files):

    total = 0

    for file in files:

        if os.path.exists(file):

            total += os.path.getsize(file)

    return total


###########################################################
# FILE COUNT BY TYPE
###########################################################

def file_type_statistics(files):

    stats = {

        "pdf": 0,

        "doc": 0,

        "image": 0,

        "unknown": 0

    }

    for file in files:

        t = get_file_type(file)

        stats[t] = stats.get(t, 0) + 1

    return stats


###########################################################
# PROCESSING SUMMARY
###########################################################

def processing_summary(files):

    return {

        "total_files": len(files),

        "statistics": file_type_statistics(files),

        "total_size":

            readable_bytes(

                total_file_size(files)

            )

    }


###########################################################
# READABLE BYTES
###########################################################

def readable_bytes(size):

    units = [

        "B",

        "KB",

        "MB",

        "GB",

        "TB"

    ]

    index = 0

    while size >= 1024 and index < len(units)-1:

        size /= 1024

        index += 1

    return f"{size:.2f} {units[index]}"


###########################################################
# OUTPUT FILE NAME
###########################################################

def output_filename(prefix="Merged_PDF"):

    now = datetime.now().strftime(

        "%Y%m%d_%H%M%S"

    )

    return f"{prefix}_{now}.pdf"


###########################################################
# OUTPUT PATH
###########################################################

def output_path():

    return os.path.join(

        OUTPUT_FOLDER,

        output_filename()

    )


###########################################################
# BUILD MERGED PDF
###########################################################

def build_final_pdf(files):

    files = remove_duplicate_files(files)

    files = sort_files(files)

    output = output_path()

    return merge_all_files(

        files,

        output

    )


###########################################################
# TEMP FILES
###########################################################

def temp_files():

    if not os.path.exists(TEMP_FOLDER):

        return []

    return [

        os.path.join(

            TEMP_FOLDER,

            f

        )

        for f in os.listdir(TEMP_FOLDER)

    ]


###########################################################
# OUTPUT FILES
###########################################################

def output_files():

    if not os.path.exists(OUTPUT_FOLDER):

        return []

    return [

        os.path.join(

            OUTPUT_FOLDER,

            f

        )

        for f in os.listdir(OUTPUT_FOLDER)

    ]


###########################################################
# DELETE EVERYTHING
###########################################################

def clean_everything():

    clean_temp_folder()

    clean_output_folder()


###########################################################
# HEALTH CHECK
###########################################################

def health_check():

    return {

        "temp_exists":

            os.path.exists(TEMP_FOLDER),

        "output_exists":

            os.path.exists(OUTPUT_FOLDER),

        "temp_files":

            len(temp_files()),

        "output_files":

            len(output_files())

    }


###########################################################
# VERSION
###########################################################

PDF_UTILS_VERSION = "1.0.0"
